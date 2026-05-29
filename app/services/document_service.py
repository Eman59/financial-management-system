import os
import uuid
from typing import Optional
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, and_
from fastapi import UploadFile

from app.config import get_settings
from app.models.document import Document
from app.schemas.document import DocumentSearchQuery

settings = get_settings()


async def extract_text_from_file(file_path: str) -> str:
    """Extract text content from uploaded file."""
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        try:
            from pypdf import PdfReader

            reader = PdfReader(file_path)
            text = ""
            for page in reader.pages:
                text += page.extract_text() or ""
            return text
        except Exception:
            return ""
    elif ext == ".docx":
        try:
            from docx import Document as DocxDocument

            doc = DocxDocument(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            return text
        except Exception:
            return ""
    elif ext == ".txt":
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    else:
        return ""


async def save_upload_file(file: UploadFile) -> str:
    """Save uploaded file to disk and return the file path."""
    os.makedirs(settings.UPLOAD_DIR, exist_ok=True)
    file_ext = os.path.splitext(file.filename)[1] if file.filename else ".bin"
    file_name = f"{uuid.uuid4()}{file_ext}"
    file_path = os.path.join(settings.UPLOAD_DIR, file_name)

    content = await file.read()
    with open(file_path, "wb") as f:
        f.write(content)

    return file_path


async def create_document(
    db: AsyncSession,
    title: str,
    company_name: str,
    document_type: str,
    file_path: str,
    uploaded_by: str,
    content: str = "",
) -> Document:
    """Create a new document record."""
    doc = Document(
        id=str(uuid.uuid4()),
        title=title,
        company_name=company_name,
        document_type=document_type,
        file_path=file_path,
        content=content,
        uploaded_by=uploaded_by,
    )
    db.add(doc)
    await db.flush()
    return doc


async def get_all_documents(db: AsyncSession, skip: int = 0, limit: int = 20) -> list[Document]:
    """Retrieve all documents with pagination."""
    result = await db.execute(
        select(Document).order_by(Document.created_at.desc()).offset(skip).limit(limit)
    )
    return list(result.scalars().all())


async def get_document_by_id(db: AsyncSession, document_id: str) -> Optional[Document]:
    """Retrieve a document by ID."""
    result = await db.execute(select(Document).where(Document.id == document_id))
    return result.scalar_one_or_none()


async def delete_document(db: AsyncSession, document_id: str) -> bool:
    """Delete a document by ID."""
    doc = await get_document_by_id(db, document_id)
    if doc is None:
        return False

    # Remove file from disk
    if os.path.exists(doc.file_path):
        os.remove(doc.file_path)

    await db.delete(doc)
    await db.flush()
    return True


async def search_documents(
    db: AsyncSession, query: DocumentSearchQuery
) -> list[Document]:
    """Search documents by metadata filters."""
    conditions = []

    if query.title:
        conditions.append(Document.title.ilike(f"%{query.title}%"))
    if query.company_name:
        conditions.append(Document.company_name.ilike(f"%{query.company_name}%"))
    if query.document_type:
        conditions.append(Document.document_type == query.document_type)
    if query.uploaded_by:
        conditions.append(Document.uploaded_by == query.uploaded_by)

    stmt = select(Document)
    if conditions:
        stmt = stmt.where(and_(*conditions))
    stmt = stmt.order_by(Document.created_at.desc())

    result = await db.execute(stmt)
    return list(result.scalars().all())
